# -*- coding: utf-8 -*-
"""生成历史事件导出模板（Word / Excel），与 preview-app 当前字段实现一致。"""
from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

ROOT = Path(__file__).resolve().parents[2]
ARTIFACT_DIR = ROOT / "artifacts" / "history-event-export-template" / "outputs"
PUBLIC_DIR = ROOT / "public" / "export-templates"
DESKTOP_DOCX = Path(r"d:\MyFiles\Desktop\历史事件.docx")
DESKTOP_XLSX = Path(r"d:\MyFiles\Desktop\历史事件.xlsx")

EXPORT_HEADERS = [
    "目标ID",
    "名单类型",
    "目标类型",
    "发现时间",
    "处置时间",
    "飞手位置",
    "目标型号",
    "识别码",
    "威胁等级",
    "验证方式",
    "探测设备",
    "反制设备",
    "处置状态",
    "备注",
]

SAMPLE_ROW = [
    "TG-2024-0001",
    "黑名单",
    "黑飞无人机",
    "2024-03-04 08:00:00",
    "2024-03-04 08:00:25",
    "未定位",
    "DJI Mavic 3",
    "1581F4100000",
    "高危",
    "人工核查",
    "雷达-01 (2.4G)",
    "干扰-01",
    "已结束",
    "多源轨迹已合并",
]

FIELD_ROWS = [
    ("目标ID", "TG-2024-0001", "融合目标编号，与列表页目标ID一致"),
    ("名单类型", "黑名单 / 白名单 / 未知", "名单归属，与列表页名单类型一致"),
    ("目标类型", "黑飞无人机 / 飞鸟 / 其他", "历史事件目标类型，与列表页一致"),
    ("发现时间", "2024-03-04 08:00:00", "事件首次发现时间"),
    ("处置时间", "2024-03-04 08:00:25 / --", "未完成处置时导出为 --"),
    ("飞手位置", "未定位", "当前实现统一展示未定位"),
    ("目标型号", "DJI Mavic 3", "识别到的目标型号"),
    ("识别码", "1581F4100000 / 未解析", "无人机识别码；无法识别时导出未解析"),
    ("威胁等级", "高危 / 中危 / 低危 / 无危", "威胁评估结果，与列表展示一致"),
    ("验证方式", "自动识别 / 人工核查", "自动识别或人工核查"),
    ("探测设备", "雷达-01 (2.4G)", "参与探测的设备"),
    ("反制设备", "干扰-01 / --", "参与反制的设备名称，未执行反制时为 --"),
    ("处置状态", "进行中 / 已结束", "事件处置状态，与列表展示一致"),
    ("备注", "多源轨迹已合并", "补充说明；等待值守类提示不导出"),
]

COL_WIDTHS = [14, 12, 14, 20, 20, 12, 16, 16, 10, 12, 18, 18, 12, 20]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: float = 10.5,
    color: str = "1F2937",
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build_docx(path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    style = document.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("历史事件数据")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = "宋体"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    meta = document.add_paragraph()
    meta_run = meta.add_run(
        "导出范围：历史事件列表当前筛选结果；如无筛选，默认导出全部数据。"
        "实际导出时按所选范围写入下方字段。"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.name = "宋体"
    meta_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    h1 = document.add_paragraph()
    h1_run = h1.add_run("一、导出字段说明")
    h1_run.bold = True
    h1_run.font.size = Pt(12)
    h1_run.font.name = "宋体"
    h1_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    field_table = document.add_table(rows=1, cols=3)
    field_table.style = "Table Grid"
    field_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, label in enumerate(["字段名", "示例值", "说明"]):
        cell = field_table.rows[0].cells[idx]
        set_cell_shading(cell, "D9E1F2")
        set_cell_text(cell, label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for name, sample, desc in FIELD_ROWS:
        cells = field_table.add_row().cells
        set_cell_text(cells[0], name, bold=True)
        set_cell_text(cells[1], sample)
        set_cell_text(cells[2], desc)

    document.add_paragraph("")
    h2 = document.add_paragraph()
    h2_run = h2.add_run("二、导出样例")
    h2_run.bold = True
    h2_run.font.size = Pt(12)
    h2_run.font.name = "宋体"
    h2_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    sample_headers = EXPORT_HEADERS[:6]
    sample_values = SAMPLE_ROW[:6]
    sample_table = document.add_table(rows=2, cols=len(sample_headers))
    sample_table.style = "Table Grid"
    for idx, text in enumerate(sample_headers):
        cell = sample_table.rows[0].cells[idx]
        set_cell_shading(cell, "E2E8F0")
        set_cell_text(cell, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx, text in enumerate(sample_values):
        set_cell_text(sample_table.rows[1].cells[idx], text, align=WD_ALIGN_PARAGRAPH.CENTER)

    note = document.add_paragraph()
    note_run = note.add_run(
        "说明：Excel / Word 导出均包含全部字段；字段顺序与历史事件列表保持一致。"
    )
    note_run.font.size = Pt(10)
    note_run.font.name = "宋体"
    note_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    document.save(path)


def _thin_border() -> Border:
    side = Side(style="thin", color="D7DFE8")
    return Border(left=side, right=side, top=side, bottom=side)


def build_xlsx(path: Path) -> None:
    wb = Workbook()
    thin = _thin_border()
    header_fill = PatternFill("solid", fgColor="1D4ED8")
    header_font = Font(name="宋体", bold=True, color="FFFFFF", size=11)
    body_font = Font(name="宋体", size=10)
    title_font = Font(name="宋体", bold=True, size=16, color="0F172A")
    sub_font = Font(name="宋体", size=10, color="475569")

    ws = wb.active
    ws.title = "导出模板"
    ws.sheet_view.showGridLines = False

    last_col = get_column_letter(len(EXPORT_HEADERS))
    ws.merge_cells(f"A1:{last_col}1")
    ws["A1"] = "历史事件数据"
    ws["A1"].font = title_font
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.merge_cells(f"A2:{last_col}2")
    ws["A2"] = (
        "适用于历史事件模块导出。实际导出时，数据量跟随导出范围（所有数据 / 查询结果 / 所选数据）。"
    )
    ws["A2"].font = sub_font
    ws["A2"].alignment = Alignment(wrap_text=True, vertical="center")

    header_row = 4
    for col, title in enumerate(EXPORT_HEADERS, start=1):
        cell = ws.cell(row=header_row, column=col, value=title)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = thin

    for col, value in enumerate(SAMPLE_ROW, start=1):
        cell = ws.cell(row=header_row + 1, column=col, value=value)
        cell.font = body_font
        cell.alignment = Alignment(vertical="center")
        cell.border = thin

    ws.merge_cells(f"A6:{last_col}6")
    ws["A6"] = "说明：正式导出时自第 5 行起写入数据，字段顺序与历史事件列表一致。"
    ws["A6"].font = Font(name="宋体", size=9, color="64748B")
    ws["A6"].alignment = Alignment(horizontal="left", vertical="center")

    for idx, width in enumerate(COL_WIDTHS, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    ws.freeze_panes = "A5"

    field_ws = wb.create_sheet("字段说明")
    field_ws.append(["字段名", "示例值", "说明"])
    for row in FIELD_ROWS:
        field_ws.append(list(row))
    for row in field_ws.iter_rows(min_row=1, max_row=field_ws.max_row, min_col=1, max_col=3):
        for cell in row:
            cell.font = body_font if cell.row > 1 else Font(name="宋体", bold=True, color="FFFFFF")
            cell.border = thin
            if cell.row == 1:
                cell.fill = PatternFill("solid", fgColor="0F766E")
                cell.alignment = Alignment(horizontal="center", vertical="center")
    field_ws.column_dimensions["A"].width = 18
    field_ws.column_dimensions["B"].width = 24
    field_ws.column_dimensions["C"].width = 48

    guide_ws = wb.create_sheet("导出说明")
    guides = [
        "历史事件导出模板说明",
        "1. Excel 模板用于列表字段对照、导出前校验与正式数据导出。",
        "2. Word 模板用于正式报表导出，适合对外呈现、归档或人工签阅。",
        "3. 导出范围支持：所有数据、查询结果、所选数据。",
        "4. 字段与 preview-app 历史事件列表及 HistoryEventItem 数据结构一致。",
        "5. 名单类型、威胁等级、验证方式等展示文案与页面列表保持一致。",
        f"6. 模板生成时间：{datetime.now():%Y-%m-%d %H:%M:%S}",
    ]
    for idx, text in enumerate(guides, start=1):
        guide_ws.cell(row=idx, column=1, value=text).font = (
            title_font if idx == 1 else Font(name="宋体", size=10, color="334155")
        )
    guide_ws.column_dimensions["A"].width = 96

    wb.save(path)


def deploy(path_docx: Path, path_xlsx: Path) -> None:
    ARTIFACT_DIR.mkdir(parents=True, exist_ok=True)
    PUBLIC_DIR.mkdir(parents=True, exist_ok=True)
    targets = [
        ARTIFACT_DIR / "历史事件.docx",
        ARTIFACT_DIR / "历史事件.xlsx",
        PUBLIC_DIR / "历史事件.docx",
        PUBLIC_DIR / "历史事件.xlsx",
        DESKTOP_DOCX,
        DESKTOP_XLSX,
    ]
    for target in targets:
        target.parent.mkdir(parents=True, exist_ok=True)
        if target.suffix.lower() == ".docx":
            shutil.copy2(path_docx, target)
        else:
            shutil.copy2(path_xlsx, target)
        print(f"Saved: {target}")


def main() -> None:
    tmp_docx = ARTIFACT_DIR / "_build.docx"
    tmp_xlsx = ARTIFACT_DIR / "_build.xlsx"
    ARTIFACT_DIR.mkdir(parents=True, exist_ok=True)
    build_docx(tmp_docx)
    build_xlsx(tmp_xlsx)
    deploy(tmp_docx, tmp_xlsx)


if __name__ == "__main__":
    main()
