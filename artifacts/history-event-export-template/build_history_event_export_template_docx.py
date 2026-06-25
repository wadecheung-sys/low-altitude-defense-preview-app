from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent / "outputs"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_PATH = OUTPUT_DIR / "history-event-export-template.docx"


FIELDS = [
    ("目标ID", "TG-2024-0001", "融合后的历史事件唯一编号"),
    ("名单状态", "黑名单", "目标当前名单归属，可为黑名单、白名单或未知"),
    ("发现时间", "2024-03-04 08:00:00", "事件首次发现时间"),
    ("处置时间", "2024-03-04 08:00:25", "执行处置动作的时间，未处置可为空"),
    ("持续时长", "00:00:25", "事件持续时长"),
    ("飞手位置", "未定位", "飞手研判定位结果"),
    ("目标型号", "DJI Mavic 3", "识别到的型号信息"),
    ("识别码", "1581F4100000", "系统统一后的识别码字段"),
    ("威胁等级", "高", "系统威胁评估结果"),
    ("处置状态", "已处置", "支持待处置、处置中、已处置、已结束四类状态"),
    ("处置结果", "驱离成功", "最终处置或结束结果"),
    ("探测设备", "雷达-01 (2.4G)", "主要探测来源设备"),
    ("反制设备", "干扰-01 (自动)", "执行联动反制的设备，未执行可为 --"),
    ("数据来源", "雷达+无线电融合", "事件来源说明"),
    ("所在区域", "核心区", "目标所在区域"),
    ("目标位置", "E:113.39 N:23.09", "目标大致位置"),
    ("威胁识别", "人工核查-真实入侵", "自动识别或人工核查结论"),
    ("备注", "多源轨迹已合并", "补充说明"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=10.5, color="1F2937", align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


document = Document()
section = document.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.1)
section.right_margin = Cm(2.1)

style = document.styles["Normal"]
style.font.name = "Microsoft YaHei"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
style.font.size = Pt(10.5)

title = document.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("历史事件导出模板")
title_run.bold = True
title_run.font.size = Pt(18)
title_run.font.color.rgb = RGBColor.from_string("0F172A")
title_run.font.name = "Microsoft YaHei"
title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

subtitle = document.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("适用于历史事件模块导出结构确认、联调演示与正式报表整理")
subtitle_run.font.size = Pt(10.5)
subtitle_run.font.color.rgb = RGBColor.from_string("475569")
subtitle_run.font.name = "Microsoft YaHei"
subtitle_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

meta = document.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
meta_run = meta.add_run("导出范围：历史事件列表当前筛选结果；如无筛选，默认导出全部数据")
meta_run.font.size = Pt(10)
meta_run.font.color.rgb = RGBColor.from_string("334155")
meta_run.font.name = "Microsoft YaHei"
meta_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

heading_1 = document.add_paragraph()
heading_1.style = document.styles["Heading 2"]
heading_1_run = heading_1.add_run("一、导出字段说明")
heading_1_run.font.name = "Microsoft YaHei"
heading_1_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
heading_1_run.font.color.rgb = RGBColor.from_string("1D4ED8")

field_table = document.add_table(rows=1, cols=3)
field_table.alignment = WD_TABLE_ALIGNMENT.CENTER
field_table.style = "Table Grid"
field_table.autofit = False
field_table.columns[0].width = Cm(3.3)
field_table.columns[1].width = Cm(4.4)
field_table.columns[2].width = Cm(8.3)

for idx, label in enumerate(["字段名", "示例值", "说明"]):
    cell = field_table.rows[0].cells[idx]
    set_cell_shading(cell, "1D4ED8")
    set_cell_text(cell, label, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)

for field_name, sample_value, desc in FIELDS:
    row_cells = field_table.add_row().cells
    set_cell_text(row_cells[0], field_name, bold=True, color="0F172A")
    set_cell_text(row_cells[1], sample_value)
    set_cell_text(row_cells[2], desc)

document.add_paragraph("")

heading_2 = document.add_paragraph()
heading_2.style = document.styles["Heading 2"]
heading_2_run = heading_2.add_run("二、导出样例")
heading_2_run.font.name = "Microsoft YaHei"
heading_2_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
heading_2_run.font.color.rgb = RGBColor.from_string("1D4ED8")

sample_table = document.add_table(rows=2, cols=6)
sample_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sample_table.style = "Table Grid"
sample_table.autofit = False
sample_table.columns[0].width = Cm(2.2)
sample_table.columns[1].width = Cm(2.8)
sample_table.columns[2].width = Cm(3.3)
sample_table.columns[3].width = Cm(2.6)
sample_table.columns[4].width = Cm(2.8)
sample_table.columns[5].width = Cm(3.8)

sample_headers = ["目标ID", "名单状态", "发现时间", "威胁等级", "处置状态", "所在区域"]
sample_values = ["TG-2024-0001", "黑名单", "2024-03-04 08:00:00", "高", "已处置", "核心区"]

for idx, text in enumerate(sample_headers):
    cell = sample_table.rows[0].cells[idx]
    set_cell_shading(cell, "E2E8F0")
    set_cell_text(cell, text, bold=True, color="0F172A", align=WD_ALIGN_PARAGRAPH.CENTER)

for idx, text in enumerate(sample_values):
    set_cell_text(sample_table.rows[1].cells[idx], text, align=WD_ALIGN_PARAGRAPH.CENTER)

document.add_paragraph("")

heading_3 = document.add_paragraph()
heading_3.style = document.styles["Heading 2"]
heading_3_run = heading_3.add_run("三、使用说明")
heading_3_run.font.name = "Microsoft YaHei"
heading_3_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
heading_3_run.font.color.rgb = RGBColor.from_string("1D4ED8")

notes = [
    "Excel 模板适合字段对照、批量整理与导出前校验。",
    "Word 模板适合形成正式汇报材料或归档附件。",
    "处置状态中的“已结束”表示未执行反制但事件已自然结束，例如无人机自离、飞鸟滋扰等场景。",
    "如后续字段扩展，建议优先保持与历史事件列表字段顺序一致。",
]

for note in notes:
    paragraph = document.add_paragraph(style=None)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(3)
    paragraph.style = document.styles["List Bullet"]
    run = paragraph.add_run(note)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string("334155")

new_section = document.add_section(WD_SECTION.CONTINUOUS)
new_section.top_margin = Cm(2)

document.save(OUTPUT_PATH)
