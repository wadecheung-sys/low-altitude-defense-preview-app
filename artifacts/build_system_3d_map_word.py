from __future__ import annotations

import json
import math
import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "系统三维地图功能设计与实现说明-工作交接版-V1.0.md"
OUTPUT = ROOT / "系统三维地图功能设计与实现说明-工作交接版-V1.0-正式版.docx"
BUILD_DIR = ROOT / "_system_3d_map_word_build"
BUILD_DIR.mkdir(parents=True, exist_ok=True)

# standard_business_brief preset token map
PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0
MARGIN_IN = 1.0
HEADER_FOOTER_IN = 0.492
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

ASCII_FONT = "Calibri"
EAST_ASIA_FONT = "Microsoft YaHei"
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
INK = "243447"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D6DCE5"
WHITE = "FFFFFF"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    ascii_font: str = ASCII_FONT,
    east_asia_font: str = EAST_ASIA_FONT,
):
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, *, size: float, color: str, bold: bool = False):
    style.font.name = ASCII_FONT
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ASCII_FONT)
    rfonts.set(qn("w:hAnsi"), ASCII_FONT)
    rfonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)


def set_paragraph_border(paragraph, *, side: str = "bottom", color: str = BORDER, size: int = 8):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), "5")
    edge.set(qn("w:color"), color)
    pbdr.append(edge)


def set_paragraph_shading(paragraph, fill: str):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_shading(cell, fill: str):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins: dict[str, int] = CELL_MARGINS_DXA):
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.find(qn("w:tcMar"))
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for side in ("top", "start", "bottom", "end"):
        node = tcmar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcmar.append(node)
        node.set(qn("w:w"), str(margins.get(side, 0)))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, *, color: str = BORDER, size: int = 4, val: str = "single"):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), val)
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def remove_table_borders(table):
    set_table_borders(table, color=WHITE, size=0, val="nil")


def set_repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def apply_table_geometry(table, widths: list[int], *, indent: int = TABLE_INDENT_DXA):
    assert sum(widths) == CONTENT_WIDTH_DXA, (widths, sum(widths))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr

    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tblw.set(qn("w:type"), "dxa")

    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent))
    tblind.set(qn("w:type"), "dxa")

    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[idx]))
            tcw.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def display_units(text: str) -> float:
    units = 0.0
    for ch in text:
        units += 1.8 if ord(ch) > 127 else 1.0
    return units


def allocate_widths(rows: list[list[str]]) -> list[int]:
    col_count = max(len(row) for row in rows)
    max_units = []
    for col in range(col_count):
        values = [row[col] if col < len(row) else "" for row in rows]
        length = max(display_units(v) for v in values)
        max_units.append(max(8.0, min(length, 56.0)) ** 0.72)
    minimum = 1050 if col_count >= 4 else 1250 if col_count == 3 else 1600
    base = minimum * col_count
    remainder = CONTENT_WIDTH_DXA - base
    weight_sum = sum(max_units)
    widths = [minimum + round(remainder * w / weight_sum) for w in max_units]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def setup_styles(doc: Document):
    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, color=INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, size=16, color=BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.keep_together = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, size=13, color=BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.keep_together = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, size=12, color=DARK_BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True
    h3.paragraph_format.keep_together = True

    if "Figure Caption" not in doc.styles:
        caption = doc.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = doc.styles["Figure Caption"]
    set_style_font(caption, size=9, color=MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = False

    if "TOC Entry" not in doc.styles:
        toc = doc.styles.add_style("TOC Entry", WD_STYLE_TYPE.PARAGRAPH)
    else:
        toc = doc.styles["TOC Entry"]
    set_style_font(toc, size=10.5, color=INK)
    toc.paragraph_format.space_before = Pt(0)
    toc.paragraph_format.space_after = Pt(3)
    toc.paragraph_format.line_spacing = 1.08


def create_numbering(doc: Document):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1

    def abstract_number(abstract_id: int, fmt: str, text: str):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        numfmt = OxmlElement("w:numFmt")
        numfmt.set(qn("w:val"), fmt)
        lvltext = OxmlElement("w:lvlText")
        lvltext.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.extend([tabs, ind, spacing])
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), ASCII_FONT)
        rfonts.set(qn("w:hAnsi"), ASCII_FONT)
        rfonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), DARK_BLUE)
        rpr.extend([rfonts, color])
        lvl.extend([start, numfmt, lvltext, suff, ppr, rpr])
        abstract.append(lvl)
        numbering.append(abstract)

    abstract_number(next_abs, "bullet", "•")
    bullet_abs = next_abs
    abstract_number(next_abs + 1, "decimal", "%1.")
    decimal_abs = next_abs + 1

    def create_num(abstract_id: int) -> int:
        nonlocal next_num
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(next_num))
        aid = OxmlElement("w:abstractNumId")
        aid.set(qn("w:val"), str(abstract_id))
        num.append(aid)
        numbering.append(num)
        value = next_num
        next_num += 1
        return value

    bullet_num = create_num(bullet_abs)

    def new_decimal_num() -> int:
        return create_num(decimal_abs)

    return bullet_num, new_decimal_num


def apply_num(paragraph, num_id: int):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, numid])
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline(paragraph, text: str, *, size: float | None = None, color: str = INK, bold: bool = False):
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=size, color=color, bold=bold)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(
                run,
                size=(size - 0.4) if size else 10.2,
                color=DARK_BLUE,
                bold=False,
                ascii_font="Consolas",
                east_asia_font=EAST_ASIA_FONT,
            )
            rpr = run._element.get_or_add_rPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "EEF2F6")
            rpr.append(shd)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=color, bold=bold)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def is_separator_row(line: str) -> bool:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c) for c in cells)


def add_markdown_table(doc: Document, raw_lines: list[str]):
    rows = parse_table(raw_lines)
    if len(rows) >= 2 and is_separator_row(raw_lines[1]):
        rows.pop(1)
    col_count = max(len(r) for r in rows)
    rows = [r + [""] * (col_count - len(r)) for r in rows]
    widths = allocate_widths(rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    apply_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    for ridx, values in enumerate(rows):
        for cidx, value in enumerate(values):
            cell = table.cell(ridx, cidx)
            if ridx == 0:
                set_cell_shading(cell, LIGHT_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.08
            if ridx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(p, value, size=9.3, color=DARK_BLUE, bold=True)
            else:
                is_short = display_units(value) <= 18 and col_count >= 3
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_short else WD_ALIGN_PARAGRAPH.LEFT
                add_inline(p, value, size=9.2, color=INK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def find_chinese_font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf") if bold else Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color="#718096", width=5):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    for delta in (2.6, -2.6):
        p = (
            end[0] + length * math.cos(angle + delta),
            end[1] + length * math.sin(angle + delta),
        )
        draw.line([end, p], fill=color, width=width)


def wrap_text_pixels(draw, text: str, font, max_width: int) -> str:
    lines: list[str] = []
    for source_line in text.splitlines() or [text]:
        current = ""
        for ch in source_line:
            candidate = current + ch
            if current and draw.textlength(candidate, font=font) > max_width:
                lines.append(current.rstrip())
                current = ch.lstrip()
            else:
                current = candidate
        if current:
            lines.append(current.rstrip())
    return "\n".join(lines)


def draw_box(draw, xy, text, *, fill, outline="#B8C4D2", text_color="#17324D", font=None, radius=18):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    box_w = x2 - x1
    wrapped = wrap_text_pixels(draw, text, font, box_w - 38)
    bbox = draw.multiline_textbbox((0, 0), wrapped, font=font, spacing=6, align="center")
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.multiline_text(
        ((x1 + x2 - tw) / 2, (y1 + y2 - th) / 2),
        wrapped,
        fill=text_color,
        font=font,
        spacing=6,
        align="center",
    )


def build_figures():
    font = find_chinese_font(34, bold=True)
    small = find_chinese_font(27, bold=False)

    # Figure 1: module relationship and business loop
    img = Image.new("RGB", (1900, 980), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 35), "三维地图模块关系与业务闭环", font=find_chinese_font(44, True), fill="#0B2545")

    inputs = [
        ("区域管理", 140),
        ("设备管理", 310),
        ("多源探测数据", 480),
        ("黑白名单主数据", 650),
    ]
    for label, y in inputs:
        draw_box(d, (70, y, 390, y + 110), label, fill="#F4F6F9", font=font)

    processes = [
        ("目标融合与航迹", 210),
        ("区域命中与身份判定", 405),
        ("威胁评估与预案匹配", 600),
    ]
    for label, y in processes:
        draw_box(d, (560, y, 930, y + 125), label, fill="#E8EEF5", outline="#88A5C2", font=font)

    draw_box(d, (1080, 330, 1450, 540), "指挥控制中心\n监测 · 告警 · 视频\n设备 · 处置", fill="#1F4D78", outline="#1F4D78", text_color="white", font=font)
    draw_box(d, (1580, 250, 1830, 390), "事件归档", fill="#E8EEF5", font=font)
    draw_box(d, (1545, 500, 1870, 635), "三维轨迹回放", fill="#F4F6F9", font=font)
    draw_box(d, (1545, 700, 1870, 835), "消息与审计", fill="#F4F6F9", font=font)

    for _, y in inputs:
        draw_arrow(d, (390, y + 55), (560, 455 if y < 400 else 665 if y > 570 else 270), width=4)
    draw_arrow(d, (930, 272), (1080, 385), width=5)
    draw_arrow(d, (930, 467), (1080, 435), width=5)
    draw_arrow(d, (930, 662), (1080, 490), width=5)
    draw_arrow(d, (1450, 390), (1580, 320), width=5)
    draw_arrow(d, (1705, 390), (1705, 500), width=5)
    draw_arrow(d, (1705, 635), (1705, 700), width=5)

    d.text((70, 880), "统一对象标识：融合目标 ID · 飞行编号 · 事件 ID · 身份 ID", font=small, fill="#667085")
    fig1 = BUILD_DIR / "figure-1-business-loop.png"
    img.save(fig1, quality=95)

    # Figure 2: real-time disposal flow
    img2 = Image.new("RGB", (1900, 980), "white")
    d2 = ImageDraw.Draw(img2)
    d2.text((70, 35), "实时监测与快速处置流程", font=find_chinese_font(44, True), fill="#0B2545")
    boxes = [
        ((70, 170, 520, 330), "1  多设备探测与目标融合"),
        ((725, 170, 1175, 330), "2  区域、名单与威胁评估"),
        ((1380, 170, 1830, 330), "3  告警展示与预案建议"),
        ((1380, 570, 1830, 730), "4  用户选中目标并选择处置方式"),
        ((725, 570, 1175, 730), "5  权限校验并匹配最近可用设备"),
        ((70, 570, 520, 730), "6  指令执行、回执、效果展示与归档"),
    ]
    fills = ["#F4F6F9", "#E8EEF5", "#DCE8F3", "#FFF4E5", "#E8EEF5", "#E7F2EE"]
    for (xy, label), fill in zip(boxes, fills):
        draw_box(d2, xy, label, fill=fill, font=font)
    draw_arrow(d2, (520, 250), (725, 250))
    draw_arrow(d2, (1175, 250), (1380, 250))
    draw_arrow(d2, (1605, 330), (1605, 570))
    draw_arrow(d2, (1380, 650), (1175, 650))
    draw_arrow(d2, (725, 650), (520, 650))
    d2.text((70, 820), "执行状态：校验中 → 等待用户确认 → 发送中 → 设备已接收 → 执行中 → 成功 / 失败 / 超时 / 已解除", font=small, fill="#667085")
    fig2 = BUILD_DIR / "figure-2-disposal-flow.png"
    img2.save(fig2, quality=95)
    return [fig1, fig2]


def set_picture_alt(picture, title: str, description: str):
    inline = picture._inline
    docpr = inline.docPr
    docpr.set("title", title)
    docpr.set("descr", description)


def setup_header_footer(doc: Document):
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer
    first_footer.paragraphs[0].text = ""

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = hp.add_run("低空防御指挥控制平台")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    hp.add_run("\t")
    right = hp.add_run("系统三维地图功能设计与实现说明")
    set_run_font(right, size=8.5, color=MUTED)
    set_paragraph_border(hp, side="bottom", color="E4E7EC", size=4)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = fp.add_run("产品工作交接版 · V1.0")
    set_run_font(left, size=8.5, color=MUTED)
    fp.add_run("\t第 ")
    add_page_field(fp)
    run = fp.add_run(" 页")
    set_run_font(run, size=8.5, color=MUTED)


def add_cover(doc: Document):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(18)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(10)
    run = kicker.add_run("产品工作交接文件")
    set_run_font(run, size=10.5, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("低空防御指挥控制平台")
    set_run_font(run, size=29, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(20)
    run = subtitle.add_run("系统三维地图功能设计与实现说明")
    set_run_font(run, size=18, color=DARK_BLUE, bold=True)

    scope = doc.add_paragraph()
    scope.paragraph_format.space_after = Pt(24)
    run = scope.add_run("指挥控制中心 · 三维轨迹 · 区域管理 · 黑白名单 · 设备与处置")
    set_run_font(run, size=11.5, color=MUTED)

    metadata = [
        ("文档版本", "V1.0"),
        ("文档状态", "产品工作交接版"),
        ("编制日期", "2026-07-02"),
        ("适用对象", "产品、三维研发、前后端研发、设备接入、测试与现场实施团队"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"{label}：")
        set_run_font(r1, size=10.5, color=DARK_BLUE, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=10.5, color=INK)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after = Pt(22)
    set_paragraph_border(rule, side="bottom", color=BLUE, size=12)

    callout = doc.add_paragraph()
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.18)
    callout.paragraph_format.space_before = Pt(0)
    callout.paragraph_format.space_after = Pt(0)
    callout.paragraph_format.line_spacing = 1.12
    set_paragraph_shading(callout, CALLOUT_FILL)
    r = callout.add_run("产品基线  ")
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    r = callout.add_run("三维地图作为统一空间业务底座，贯通区域规划、目标融合、威胁评估、实时处置与历史复盘。")
    set_run_font(r, size=10.5, color=INK)

    doc.add_page_break()


def add_toc(doc: Document, lines: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("目录")
    set_run_font(r, size=20, color=NAVY, bold=True)
    set_paragraph_border(p, side="bottom", color="D6DCE5", size=6)
    for line in lines:
        match = re.match(r"^##\s+(.+)$", line)
        if not match:
            continue
        p = doc.add_paragraph(style="TOC Entry")
        add_inline(p, match.group(1), size=10.5, color=INK)
    doc.add_page_break()


def add_figure(doc: Document, path: Path, caption: str, description: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    picture = run.add_picture(str(path), width=Inches(6.35))
    set_picture_alt(picture, caption, description)
    cp = doc.add_paragraph(style="Figure Caption")
    cp.add_run(caption)


def build_document():
    source_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    content_start = next(i for i, line in enumerate(source_lines) if line.startswith("## 1."))
    body_lines = source_lines[content_start:]
    figures = build_figures()

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.top_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.header_distance = Inches(HEADER_FOOTER_IN)
    section.footer_distance = Inches(HEADER_FOOTER_IN)

    setup_styles(doc)
    bullet_num, new_decimal_num = create_numbering(doc)
    setup_header_footer(doc)
    add_cover(doc)
    add_toc(doc, body_lines)

    current_decimal_num = None
    in_decimal_list = False
    figure_index = 0
    i = 0
    while i < len(body_lines):
        line = body_lines[i]
        stripped = line.strip()

        if not stripped:
            in_decimal_list = False
            current_decimal_num = None
            i += 1
            continue

        if stripped.startswith("```mermaid"):
            i += 1
            while i < len(body_lines) and not body_lines[i].strip().startswith("```"):
                i += 1
            if figure_index == 0:
                add_figure(
                    doc,
                    figures[0],
                    "图 1  三维地图模块关系与业务闭环",
                    "区域、设备、探测数据和名单经过融合、区域判定、威胁评估与预案匹配进入指挥控制中心，最终归档并用于三维轨迹回放和审计。",
                )
            else:
                add_figure(
                    doc,
                    figures[1],
                    "图 2  实时监测与快速处置流程",
                    "从多设备探测、目标融合、威胁评估到用户选择处置方式、权限校验、就近设备执行与事件归档的六步流程。",
                )
            figure_index += 1
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(body_lines) and is_separator_row(body_lines[i + 1]):
            table_lines = [line, body_lines[i + 1]]
            i += 2
            while i < len(body_lines) and body_lines[i].strip().startswith("|"):
                table_lines.append(body_lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            in_decimal_list = False
            current_decimal_num = None
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1)) - 1
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, heading.group(2), color=BLUE if level < 3 else DARK_BLUE, bold=True)
            i += 1
            in_decimal_list = False
            current_decimal_num = None
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            p = doc.add_paragraph()
            apply_num(p, bullet_num)
            add_inline(p, bullet.group(1), size=11, color=INK)
            i += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            if not in_decimal_list or current_decimal_num is None:
                current_decimal_num = new_decimal_num()
                in_decimal_list = True
            p = doc.add_paragraph()
            apply_num(p, current_decimal_num)
            add_inline(p, numbered.group(1), size=11, color=INK)
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, stripped, size=11, color=INK)
        i += 1
        in_decimal_list = False
        current_decimal_num = None

    doc.core_properties.title = "低空防御指挥控制平台——系统三维地图功能设计与实现说明"
    doc.core_properties.subject = "产品工作交接版"
    doc.core_properties.author = "低空防御指挥控制平台项目组"
    doc.core_properties.keywords = "三维地图, 指挥控制中心, 区域管理, 飞行轨迹, 黑白名单"
    doc.core_properties.comments = ""

    # Preset audit before save
    assert round(section.page_width.inches, 3) == PAGE_WIDTH_IN
    assert round(section.left_margin.inches, 3) == MARGIN_IN
    assert round(section.header_distance.inches, 3) == HEADER_FOOTER_IN
    assert figure_index == 2
    for table in doc.tables:
        tblw = table._tbl.tblPr.find(qn("w:tblW"))
        assert tblw is not None and int(tblw.get(qn("w:w"))) == CONTENT_WIDTH_DXA
        grid_widths = [int(x.get(qn("w:w"))) for x in table._tbl.tblGrid.findall(qn("w:gridCol"))]
        assert sum(grid_widths) == CONTENT_WIDTH_DXA

    doc.save(OUTPUT)
    report = {
        "output": str(OUTPUT),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "figures": figure_index,
        "preset": "standard_business_brief",
        "header_pattern": "memo_masthead",
        "table_width_dxa": CONTENT_WIDTH_DXA,
    }
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    build_document()
