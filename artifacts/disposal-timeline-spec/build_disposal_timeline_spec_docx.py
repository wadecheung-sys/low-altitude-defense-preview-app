from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "outputs"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_PATH = OUTPUT_DIR / "设备处置时间链实现说明-正式交接版.docx"

FONT = "Microsoft YaHei"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GREEN = "DFF3E8"
YELLOW = "FFF3CD"
RED = "FDE2E2"


def set_east_asia_font(run, font=FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=INK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_east_asia_font(run)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D7DFE8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(0, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])


def add_heading(document, text, level=1):
    p = document.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    set_east_asia_font(run)
    run.font.color.rgb = RGBColor.from_string(BLUE if level < 3 else DARK_BLUE)
    return p


def add_body(document, text):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(INK)
    return p


def add_bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.95)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(10.2)
    run.font.color.rgb = RGBColor.from_string(INK)
    return p


def add_code_block(document, text):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    set_table_borders(table, "CBD5E1")
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F8FAFC")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(8.8)
    run.font.color.rgb = RGBColor.from_string("334155")
    document.add_paragraph("")


def add_table(document, headers, rows, widths, header_fill=LIGHT_BLUE, font_size=9.2):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        set_cell_text(cell, header, bold=True, color="0F172A", size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=font_size)
    document.add_paragraph("")
    return table


document = Document()
section = document.sections[0]
section.orientation = WD_ORIENTATION.PORTRAIT
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)
section.header_distance = Cm(1.25)
section.footer_distance = Cm(1.25)

styles = document.styles
styles["Normal"].font.name = FONT
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
styles["Normal"].font.size = Pt(10.5)

for style_name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
]:
    style = styles[style_name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.25

title = document.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(3)
run = title.add_run("设备处置时间链实现说明")
set_east_asia_font(run)
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor.from_string("0F172A")

subtitle = document.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(10)
run = subtitle.add_run("飞行记录详情页 - 阶段语义、展示形态、状态规则与研发交接规范")
set_east_asia_font(run)
run.font.size = Pt(10.5)
run.font.color.rgb = RGBColor.from_string(MUTED)

meta_table = add_table(
    document,
    ["项目", "说明"],
    [
        ["适用模块", "历史事件 / 飞行记录详情 / 设备处置时间链"],
        ["文档用途", "用于正式工作交接、前端实现、脚本文案生成与演示数据维护"],
        ["阶段范围", "设备发现、威胁识别、威胁评估、处置执行、处置结果"],
        ["核心原则", "五个阶段作为标准模型保留；页面只展示已发生、正在发生或具有解释价值的未触发阶段"],
        ["模块提示文案", "展示小型低空飞行物从发现、识别、评估、处置到复盘归档的全过程记录。"],
    ],
    [2100, 7260],
    header_fill=LIGHT_GRAY,
)

add_heading(document, "一、设计目标与边界", 1)
add_body(
    document,
    "设备处置时间链用于解释一条飞行记录从原始发现到最终闭环的过程依据。上方基本信息展示事件当前快照，时间链展示该快照形成的过程原因与阶段结论。"
)
add_bullet(document, "固定阶段：阶段名称与顺序固定为设备发现、威胁识别、威胁评估、处置执行、处置结果。")
add_bullet(document, "展示边界：不展示空节点；未开始阶段默认不展示；未触发阶段仅在具有业务解释价值时展示。")
add_bullet(document, "文案原则：所有提示文案由标准模板与结构化变量拼接生成，避免自由口语化描述。")
add_bullet(document, "冗余控制：时间链允许保留关键结论锚点，但不重复上方基本信息中的全量字段。")

add_heading(document, "二、推荐呈现形式", 1)
add_body(
    document,
    "飞行记录详情页推荐采用纵向时间线。该页面信息密度较高，纵向时间线更适合承载阶段摘要、详情字段与标签。横向步骤条可作为未来列表页或详情页顶部概览，但不建议替代详情页主体时间链。"
)
add_table(
    document,
    ["设计项", "推荐方案", "说明"],
    [
        ["方向", "纵向时间线", "适合不等长内容，便于展示跳过、未触发、进行中等差异状态。"],
        ["阶段卡片", "标题 + 时间 + 状态标识 + 概要 + 折叠详情 + 业务标签", "保持每个阶段结构一致，便于脚本生成。"],
        ["折叠详情", "默认收起，展开后纵向展示字段", "详情内容不铺满整行，避免页面臃肿；字段按 label/value 纵向排列。"],
        ["状态标识", "已完成 / 进行中 / 未触发 / 未开始", "状态标识用于表达流程进度，不作为业务标签使用。"],
        ["业务标签", "设备类型、识别方式、评估方式、处置方式、处置结论", "由字段脚本派生，只展示有业务含义的短标签。"],
        ["未开始节点", "默认隐藏", "详情页面向既有事件记录，隐藏未发生内容可减少误解。"],
        ["未触发节点", "可展示为轻量节点", "例如飞鸟躁扰、白名单低风险、未命中预案等，需要解释为什么没有反制。"],
    ],
    [1700, 2300, 5360],
)

add_heading(document, "三、统一数据结构建议", 1)
add_body(
    document,
    "前端组件建议只接收结构化时间链节点，所有文案生成、状态判定、标签派生放在构建函数中完成。"
)
add_code_block(
    document,
    """type TimelineStageStatus = 'done' | 'current' | 'skipped' | 'pending'

type TimelineStage = {
  stageCode: 'discovery' | 'identification' | 'assessment' | 'disposal' | 'result'
  stageName: string
  timestamp: string | null
  status: TimelineStageStatus
  summary: string
  details: Array<{ label: string; value: string }>
  tags: string[]
  source: string
  visible: boolean
}""",
)

add_heading(document, "四、阶段语义与展示字段", 1)
add_table(
    document,
    ["阶段", "业务含义", "推荐展示字段", "概要模板"],
    [
        [
            "设备发现",
            "展示原始探测事实，不输出身份判断。",
            "发现时间、发现设备、设备类型、原始位置、频段/频率、速度/高度、数据来源。",
            "[{time}] {deviceNames}发现疑似低空目标，已形成原始探测记录。",
        ],
        [
            "威胁识别",
            "基于多源融合识别目标身份和类型。",
            "目标类型、合作属性、名单属性、品牌型号、识别码、置信度、人工核查结论。",
            "[{time}] 完成目标识别，判定为{targetType}，身份属性为{identityLabel}。",
        ],
        [
            "威胁评估",
            "根据威胁评估模块规则输出威胁等级。",
            "威胁等级、命中规则、评估依据、所在区域、规则版本。",
            "[{time}] 根据威胁评估规则，目标命中{ruleCount}项条件，评估结果为{threatLevel}。",
        ],
        [
            "处置执行",
            "展示预案策略触发和反制设备动作，不承担最终结果。",
            "策略名称、触发条件、联动设备、执行动作、执行方式、指令状态。",
            "[{time}] 命中预案策略“{planName}”，已联动{deviceCount}台反制设备执行{actionName}。",
        ],
        [
            "处置结果",
            "展示事件闭环结论与最终总结。",
            "处置状态、最终结果、结束原因、人工核查、结果说明。",
            "[{time}] 事件已闭环，最终结果为{resultLabel}。",
        ],
    ],
    [1500, 2350, 2750, 2760],
    font_size=8.8,
)

add_heading(document, "五、状态与可见性规则", 1)
add_table(
    document,
    ["状态", "页面文案", "触发条件", "展示规则"],
    [
        ["done", "已完成", "阶段已发生并产生结构化结论。", "正常展示完整节点。"],
        ["current", "进行中", "前置阶段已完成，本阶段正在执行或等待回传。", "进行中事件停留在处置执行阶段；处置完成后再展示处置结果。"],
        ["skipped", "未触发", "阶段不适用或未命中策略，但该事实需要解释。", "展示轻量节点，摘要必须说明未触发原因。"],
        ["pending", "未开始", "前置阶段尚未完成，阶段未发生。", "详情页默认隐藏，不作为空节点展示。"],
    ],
    [1500, 1500, 3500, 2860],
)
add_body(
    document,
    "处置执行阶段是否展示取决于业务解释价值：已触发反制时完整展示；未命中预案但需要说明原因时展示为“未触发”；仅因流程未走到此处时隐藏。进行中事件保留在处置执行阶段，阶段完成后再展示处置结果。"
)

add_heading(document, "六、提示文案构成规则", 1)
add_body(
    document,
    "提示文案由“固定句式 + 可替换变量”组成。变量必须来自事件记录、设备记录、识别结果、评估规则或预案策略，不建议在页面组件中临时拼接自由文案。"
)
add_table(
    document,
    ["文案类型", "固定构成", "可替换变量"],
    [
        ["阶段概要", "时间 + 阶段动作 + 核心结论", "{time}、{deviceNames}、{targetType}、{threatLevel}、{resultLabel}"],
        ["阶段详情", "字段名 + 字段值", "{frequency}、{location}、{ruleNames}、{planName}、{actionName}"],
        ["未触发说明", "未触发原因 + 后续状态", "{skipReason}、{handlingStatus}、{endReason}"],
        ["业务标签", "字段派生标签", "{deviceType}、{recognitionMode}、{targetType}、{assessmentMode}、{disposalMode}、{resultAction}"],
    ],
    [1700, 3200, 4460],
)
add_code_block(
    document,
    """const templates = {
  discovery: '[{time}] {deviceNames}发现疑似低空目标，已形成原始探测记录。',
  identification: '[{time}] 完成目标识别，判定为{targetType}，身份属性为{identityLabel}。',
  assessment: '[{time}] 根据威胁评估规则，目标命中{ruleCount}项条件，评估结果为{threatLevel}。',
  disposal: '[{time}] 命中预案策略“{planName}”，已联动{deviceCount}台反制设备执行{actionName}。',
  disposalSkipped: '[{time}] 未命中反制预案，未下发处置指令；原因：{skipReason}。',
  result: '[{time}] 事件已闭环，最终结果为{resultLabel}。'
}""",
)

add_heading(document, "七、业务标签派生规则", 1)
add_table(
    document,
    ["阶段", "标签内容", "派生条件", "不展示内容"],
    [
        [
            "设备发现",
            "雷达、无线电、光电、多源融合",
            "来自发现设备、数据来源或融合来源字段，表示参与发现和分析运算的设备类型。",
            "不展示具体设备名、设备编号或完整数据来源句子。",
        ],
        [
            "威胁识别",
            "系统识别、人工识别、无人机、疑似无人机、飞鸟躁扰、地面杂波、黑名单、白名单、非合作式",
            "来自识别方式、目标类型、名单属性与合作属性。",
            "不使用“非无人机目标”等泛化标签；应展示具体识别类型。",
        ],
        [
            "威胁评估",
            "系统评估、人工核查",
            "来自威胁评估结论的产生方式。系统自动计算时为系统评估，人工介入复核时为人工核查。",
            "不展示威胁高、威胁中、威胁低等阶段表现标签；威胁等级放入概要或详情。",
        ],
        [
            "处置执行",
            "预案联动、人工处置",
            "来自处置动作来源。命中预案策略并联动设备时为预案联动，人工操作或人工确认触发时为人工处置。",
            "不展示指令已下发、指令执行中等流程状态标签；流程状态由阶段状态标识表达。",
        ],
        [
            "处置结果",
            "迫降、驱离、自离、压制、打击、扰动排除、误报排除",
            "来自最终处置结果或结束原因，用具体结论表达事件闭环结果。",
            "不展示已处置、已结束等状态类标签；状态由阶段状态标识表达。",
        ],
    ],
    [1400, 2400, 3300, 2260],
    font_size=8.5,
)

add_heading(document, "八、常见场景展示状态", 1)
add_body(
    document,
    "以下场景用于演示和研发联调。实际展示时，字段可按事件数据替换；阶段状态与可见性应保持一致。"
)

scenario_rows = [
    [
        "黑飞无人机入侵",
        "非合作式无人机进入防护区域。",
        "设备发现=已完成；威胁识别=已完成；威胁评估=已完成；处置执行=已完成/进行中；处置结果=已完成或暂不展示。",
        "已闭环事件展示完整五阶段；处置进行中事件停留在处置执行阶段，处置结果暂不展示。",
    ],
    [
        "黑名单无人机入侵",
        "合作式或品牌无人机命中黑名单后进入区域。",
        "设备发现=已完成；威胁识别=已完成；威胁评估=已完成；处置执行=已完成；处置结果=已完成。",
        "展示完整五阶段，识别阶段突出黑名单命中，评估阶段突出命中规则。",
    ],
    [
        "白名单无人机",
        "合作式白名单目标进入授权或低风险区域。",
        "设备发现=已完成；威胁识别=已完成；威胁评估=已完成；处置执行=未触发；处置结果=已完成。",
        "白名单目标同样经过威胁评估；处置执行展示为轻量未触发节点，说明未命中反制预案原因。",
    ],
    [
        "飞鸟躁扰",
        "探测信号经识别判定为飞鸟躁扰。",
        "设备发现=已完成；威胁识别=已完成；威胁评估=已完成；处置执行=未触发；处置结果=已完成。",
        "飞鸟躁扰归入已结束；处置结果展示为扰动排除或误报排除。",
    ],
]
add_table(
    document,
    ["场景", "业务描述", "阶段状态", "展示说明"],
    scenario_rows,
    [1500, 2300, 3300, 2260],
    font_size=8.5,
)

add_heading(document, "九、场景样例文案", 1)
add_table(
    document,
    ["场景", "阶段", "状态", "概要示例"],
    [
        ["黑飞无人机入侵", "设备发现", "已完成", "[2026-06-25 14:22:31] 雷达-01、无线电侦测-02发现疑似低空目标，已形成原始探测记录。"],
        ["黑飞无人机入侵", "威胁识别", "已完成", "[2026-06-25 14:23:08] 完成目标识别，判定为无人机，身份属性为非合作式黑飞目标。"],
        ["黑飞无人机入侵", "威胁评估", "已完成", "[2026-06-25 14:23:15] 根据威胁评估规则，目标命中2项条件，评估结果为高。"],
        ["黑飞无人机入侵", "处置执行", "已完成", "[2026-06-25 14:23:18] 命中预案策略“核心区黑飞无人机反制”，已联动1台反制设备执行定向驱离。"],
        ["黑名单无人机入侵", "威胁识别", "已完成", "[2026-06-25 15:10:08] 完成目标识别，判定为合作式无人机，身份属性为黑名单目标。"],
        ["白名单无人机", "处置执行", "未触发", "[2026-06-25 16:08:20] 未命中反制预案，未下发处置指令；原因：白名单目标且未触发异常行为条件。"],
        ["飞鸟躁扰", "威胁识别", "已完成", "[2026-06-25 17:31:11] 完成目标识别，判定为飞鸟躁扰。"],
        ["飞鸟躁扰", "处置结果", "已完成", "[2026-06-25 17:31:42] 事件已闭环，最终结果为扰动排除。"],
    ],
    [1600, 1500, 1200, 5060],
    font_size=8.3,
)

add_heading(document, "十、与详情页其他区域的冗余控制", 1)
add_body(
    document,
    "基本信息区展示事件快照，时间链展示阶段过程。若两个区域出现同一字段，时间链应仅保留该字段对阶段判断有解释作用的部分。"
)
add_table(
    document,
    ["信息类型", "基本信息区", "时间链"],
    [
        ["目标ID、识别码、型号", "完整展示", "仅在设备发现或识别阶段作为辅助详情，不进入概要。"],
        ["威胁等级", "展示最终等级", "展示等级产生原因，例如命中规则和评估依据。"],
        ["反制设备", "展示最终联动设备", "展示命中策略、动作和指令状态。"],
        ["处置状态/结果", "展示当前最终状态", "展示该状态形成的原因和闭环时间。"],
        ["探测设备", "展示汇总设备", "展示发现阶段的来源设备与原始信号摘要。"],
    ],
    [2200, 3600, 3560],
)

add_heading(document, "十一、开发落地建议", 1)
add_bullet(document, "使用五个构建函数生成节点：buildDiscoveryStage、buildIdentificationStage、buildAssessmentStage、buildDisposalStage、buildResultStage。")
add_bullet(document, "组件只负责渲染，不在组件中判断复杂业务文案。")
add_bullet(document, "缺失字段统一展示为“--”，但不应把缺失字段拼入概要句。")
add_bullet(document, "详情页默认隐藏 pending 节点；skipped 节点必须提供 skipReason。")
add_bullet(document, "业务标签由字段派生，不作为人工维护的装饰文本。")

add_heading(document, "十二、交接结论", 1)
add_bullet(document, "白名单目标必须进入威胁评估阶段，威胁评估结果决定后续是否触发处置。")
add_bullet(document, "飞鸟躁扰归入已结束事件，处置结果以扰动排除、误报排除等具体结论展示。")
add_bullet(document, "进行中事件停留在处置执行阶段；处置执行完成后再展示处置结果阶段。")
add_bullet(document, "详情页采用纵向时间线；阶段详情默认折叠，展开后以纵向字段列表展示。")
add_bullet(document, "业务标签只表达业务分类、来源方式或最终结论，不表达阶段进度状态。")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run("设备处置时间链实现说明 | 历史事件模块")
set_east_asia_font(footer_run)
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor.from_string(MUTED)

document.save(OUTPUT_PATH)
print(OUTPUT_PATH)
