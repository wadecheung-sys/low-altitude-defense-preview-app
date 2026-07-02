# -*- coding: utf-8 -*-
"""生成《处置时间链场景说明》— 对齐 preview-app 最新实现（V1.0.6）。"""
from __future__ import annotations

import shutil
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "outputs"
WORKSPACE_DOC = (
    Path(__file__).resolve().parents[3]
    / "workspace"
    / "docs"
    / "处置时间链场景说明.docx"
)
PRODUCT_DOC = Path(
    r"e:\同方威视工作素材\工作热区\低空防御\3.工作对接\处置时间链场景说明.docx"
)
DESKTOP_DOC = Path(r"d:\MyFiles\Desktop\处置时间链场景说明.docx")
OUTPUT_PATH = OUTPUT_DIR / "处置时间链场景说明.docx"

FONT = "Microsoft YaHei"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"

STAGE_NAMES = ["目标发现", "威胁识别", "威胁评估", "处置执行", "目标结果"]


def set_east_asia_font(run, font=FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=INK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_east_asia_font(run)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D7DFE8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(0, grid)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])


def add_heading(document, text, level=1):
    p = document.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    set_east_asia_font(run)
    run.font.color.rgb = RGBColor.from_string(BLUE if level < 3 else DARK_BLUE)
    return p


def add_body(document, text):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(INK)
    return p


def add_bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.95)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(10.2)
    run.font.color.rgb = RGBColor.from_string(INK)
    return p


def add_table(document, headers, rows, widths, header_fill=LIGHT_BLUE, font_size=9.0):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        set_cell_text(
            cell, header, bold=True, color="0F172A", size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER
        )
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=font_size)
    document.add_paragraph("")
    return table


def add_scenario_section(document, title, trigger, flow_summary, stage_rows):
    add_heading(document, title, 2)
    add_body(document, f"触发条件：{trigger}")
    add_body(document, f"流程特征：{flow_summary}")
    add_table(
        document,
        ["阶段", "节点状态", "摘要文案要点", "阶段详情（可折叠）", "概要示例"],
        stage_rows,
        [1050, 950, 2100, 2100, 2660],
        font_size=8.3,
    )


SCENARIOS = [
    {
        "title": "3.1 黑飞无人机入侵",
        "trigger": (
            "historyTargetType=黑飞无人机，或识别码 uavSn=未解析；名单类型通常为「未知」；"
            "目标进入防护区域或触发异常行为规则。"
        ),
        "flow": (
            "五阶段完整展示。威胁评估可能按低危→中危→高危递进输出多条摘要；"
            "中危及以上且命中预案时，处置执行为「已完成」并展示预案联动详情；"
            "低危或未命中预案时为「未触发」，阶段详情为空，摘要说明原因。"
        ),
        "stages": [
            [
                "目标发现",
                "已完成",
                "模板：{时间}{设备类型}{设备名称}{设备ID}发现目标",
                "数据来源、发现设备、目标ID、逼近距离",
                "2025-03-01 10:00:08雷达-01DEV-0008发现目标",
            ],
            [
                "威胁识别",
                "已完成",
                "模板：{时间}根据多源数据融合成果，确认目标类型为{目标类型}",
                "识别结论（系统/人工核查）、目标类型、名单类型、目标型号、识别码",
                "2025-03-01 10:00:16根据多源数据融合成果，确认目标类型为黑飞无人机",
            ],
            [
                "威胁评估",
                "已完成",
                "按递进路径输出 1～3 条评估摘要（低危/中危/高危）",
                "威胁等级、命中规则、评估依据、飞行状态；（多次评估时含评估次序）",
                "2025-03-01 10:00:24…判定为中危（若最终高危则先低危、再中危、再高危）",
            ],
            [
                "处置执行",
                "已完成/进行中/未触发",
                "自动：命中预案规则…执行…；未触发：{时间}未命中反制预案，未下发处置指令。",
                "已触发且命中预案：策略名称、触发规则、天气要素、所处区域、反制动作、联动设备；"
                "未触发/未命中：无阶段详情",
                "2025-03-01 10:01:02命中预案规则「非合作式无人机入侵反制」，执行定向驱离，执行设备干扰-01",
            ],
            [
                "目标结果",
                "已完成",
                "驱离/自离：{时间}，编码为{编码}的{无人机/非无人机}目标消失,逗留时间为{逗留时间}",
                "不展示阶段详情面板",
                "2025-03-01 10:05:30，编码为T-0008的无人机目标消失,逗留时间为00:05:22",
            ],
        ],
    },
    {
        "title": "3.2 未知无人机入侵（合作式、未备案）",
        "trigger": (
            "识别码可解析（uavSn≠未解析）；名单类型=未知；historyTargetType=合作式无人机；"
            "目标进入监测/防护区域并触发威胁评估规则。"
        ),
        "flow": (
            "识别阶段突出名单类型=未知；评估使用常规无人机目标评估规则或区域入侵规则；"
            "高危时可能命中「高威胁目标联动反制」或「常规无人机处置预案」；低危可未触发反制。"
        ),
        "stages": [
            [
                "目标发现",
                "已完成",
                "同目标发现模板",
                "数据来源、发现设备、目标ID、逼近距离",
                "2025-03-01 11:02:10无线电-02DEV-0012发现目标",
            ],
            [
                "威胁识别",
                "已完成",
                "目标类型=合作式无人机；名单类型=未知",
                "识别结论、目标类型、名单类型、目标型号、识别码（展示具体 SN）",
                "2025-03-01 11:02:18根据多源数据融合成果，确认目标类型为合作式无人机",
            ],
            [
                "威胁评估",
                "已完成",
                "规则名通常为「常规无人机目标评估规则」或区域相关规则",
                "同威胁评估阶段详情结构",
                "2025-03-01 11:02:26…判定为低危（或递进至中危/高危）",
            ],
            [
                "处置执行",
                "已完成/未触发",
                "视威胁等级与预案命中情况",
                "命中预案时展示完整联动详情；未命中时详情为空",
                "2025-03-01 11:03:10命中预案规则「常规无人机处置预案」，执行定向驱离，执行设备干扰-02",
            ],
            [
                "目标结果",
                "已完成",
                "自然飞离时为驱离/自离模板",
                "无阶段详情",
                "2025-03-01 11:08:00，编码为T-0012的无人机目标消失,逗留时间为00:04:50",
            ],
        ],
    },
    {
        "title": "3.3 黑名单无人机入侵",
        "trigger": "名单类型=黑名单；目标进入核心/管制等重点区域或触发黑名单关联规则。",
        "flow": (
            "评估规则通常为「黑名单目标进入重点区域」或「黑名单目标活动」；"
            "处置执行多命中「黑名单无人机入侵反制」预案；五阶段完整展示。"
        ),
        "stages": [
            [
                "目标发现",
                "已完成",
                "同目标发现模板",
                "数据来源、发现设备、目标ID、逼近距离",
                "2025-03-01 12:00:05雷达-01 (2.4G) DEV-0020发现目标",
            ],
            [
                "威胁识别",
                "已完成",
                "目标类型=合作式或黑飞；名单类型=黑名单",
                "识别结论含系统识别/人工核查标签",
                "2025-03-01 12:00:13根据多源数据融合成果，确认目标类型为合作式无人机",
            ],
            [
                "威胁评估",
                "已完成",
                "威胁等级通常为中危或高危；可能递进展示",
                "命中规则=黑名单目标进入重点区域/活动",
                "2025-03-01 12:00:21…判定为高危",
            ],
            [
                "处置执行",
                "已完成",
                "策略=黑名单无人机入侵反制",
                "策略名称、触发规则、天气、区域、反制动作、联动设备",
                "2025-03-01 12:01:00命中预案规则「黑名单无人机入侵反制」，执行定向驱离，执行设备干扰-01",
            ],
            [
                "目标结果",
                "已完成",
                "结果可为驱离、迫降、压制、打击等",
                "无阶段详情；迫降/打击使用对应专用模板",
                "2025-03-01 12:06:00，编码为T-0020的无人机目标消失,逗留时间为00:05:55",
            ],
        ],
    },
    {
        "title": "3.4 白名单无人机飞行（仅监测）",
        "trigger": (
            "名单类型=白名单；目标在授权区域通行或未触发异常行为；"
            "威胁等级通常为低危/无危。"
        ),
        "flow": (
            "发现、识别、评估正常展示；评估规则为「白名单目标授权通行」；"
            "处置执行节点状态=未触发，摘要含 skipReason「白名单目标未触发异常行为条件」，阶段详情为空；"
            "目标结果以监测闭环归档。"
        ),
        "stages": [
            [
                "目标发现",
                "已完成",
                "同目标发现模板",
                "数据来源、发现设备、目标ID、逼近距离",
                "2025-03-01 13:00:00光电-01DEV-0030发现目标",
            ],
            [
                "威胁识别",
                "已完成",
                "名单类型=白名单",
                "识别结论、目标类型、名单类型、目标型号、识别码",
                "2025-03-01 13:00:08根据多源数据融合成果，确认目标类型为合作式无人机",
            ],
            [
                "威胁评估",
                "已完成",
                "最终等级多为低危或仅无危/低危单步",
                "命中规则=白名单目标授权通行",
                "2025-03-01 13:00:16…判定为低危",
            ],
            [
                "处置执行",
                "未触发",
                "摘要：{时间}未命中反制预案，未下发处置指令。",
                "阶段详情为空（不展示联动设备）",
                "2025-03-01 13:00:28未命中反制预案，未下发处置指令。",
            ],
            [
                "目标结果",
                "已完成",
                "驱离/自离或监测归档",
                "无阶段详情",
                "2025-03-01 13:10:00，编码为T-0030的无人机目标消失,逗留时间为00:09:40",
            ],
        ],
    },
    {
        "title": "3.5 躁扰告警（飞鸟/杂波）",
        "trigger": (
            "historyTargetType=躁扰信号-飞鸟，或 manualConfirmStatus=躁扰告警/人工-躁扰告警；"
            "目标类型为非无人机躁扰信号。"
        ),
        "flow": (
            "威胁评估仅输出无危单步；处置执行未触发，原因=「目标识别为飞鸟或躁扰信号，未进入反制处置流程」；"
            "目标结果摘要为「事件已闭环，目标判定为躁扰信号排除」，不使用驱离/自离模板。"
        ),
        "stages": [
            [
                "目标发现",
                "已完成",
                "同目标发现模板",
                "数据来源、发现设备、目标ID、逼近距离",
                "2025-03-01 14:00:02雷达-03 (5.8G) DEV-0008发现目标",
            ],
            [
                "威胁识别",
                "已完成",
                "目标类型=躁扰信号-飞鸟；人工核查时识别结论=人工核查：躁扰告警",
                "识别结论、目标类型、名单类型、目标型号、识别码",
                "2025-03-01 14:00:10根据多源数据融合成果，确认目标类型为躁扰信号-飞鸟",
            ],
            [
                "威胁评估",
                "已完成",
                "固定无危单步；规则=目标机型排除规则",
                "威胁等级=无危；评估依据=无（任意场景）",
                "2025-03-01 14:00:18…判定为无危",
            ],
            [
                "处置执行",
                "未触发",
                "摘要含 skipReason，阶段详情为空",
                "无阶段详情",
                "2025-03-01 14:00:30未命中反制预案，未下发处置指令。",
            ],
            [
                "目标结果",
                "已完成",
                "固定文案：事件已闭环，目标判定为躁扰信号排除",
                "无阶段详情",
                "2025-03-01 14:02:00事件已闭环，目标判定为躁扰信号排除。",
            ],
        ],
    },
]


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = document.styles
    styles["Normal"].font.name = FONT
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("处置时间链场景说明")
    set_east_asia_font(run)
    run.bold = True
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("历史事件 · 飞行记录详情 · 五类典型场景阶段展示与触发流程")
    set_east_asia_font(run)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"版本：V2.0（对齐 preview-app 实现）  |  更新日期：{date.today().isoformat()}")
    set_east_asia_font(run)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    document.add_paragraph("")

    add_table(
        document,
        ["项目", "说明"],
        [
            ["适用页面", "历史事件列表 → 飞行记录详情（IncidentTargetDetail）"],
            ["布局", "左侧 50% 设备处置时间链 + 右侧 50% 无人机光电；下方地图轨迹与航机列表"],
            [
                "阶段模型",
                "目标发现 → 威胁识别 → 威胁评估 → 处置执行 → 目标结果（五阶段固定顺序）",
            ],
            [
                "展示原则",
                "仅展示 status≠pending 的节点；未开始节点隐藏；"
                "未触发处置须展示摘要原因；目标结果仅摘要、无阶段详情",
            ],
            ["废止说明", "本文档 V2.0 取代旧版「设备发现/处置结果」口径说明，以当前代码为准"],
        ],
        [2100, 7260],
        header_fill=LIGHT_GRAY,
    )

    add_heading(document, "一、阶段模型与节点状态", 1)
    add_body(
        document,
        "设备处置时间链用于解释一条历史事件从首次发现到闭环归档的过程。"
        "各场景共用同一五阶段顺序，差异体现在目标类型、名单属性、威胁递进评估、"
        "是否命中预案联动反制，以及最终目标结果文案。"
    )
    add_table(
        document,
        ["阶段", "界面标题", "业务含义", "节点状态", "阶段详情"],
        [
            ["discover", "目标发现", "记录探测事实，文案为「发现目标」", "固定 done（已完成）", "有"],
            ["threat", "威胁识别", "融合确认目标类型、名单、识别码等", "固定 done", "有"],
            ["assess", "威胁评估", "按规则输出威胁等级，可多级递进", "固定 done", "有（含多次评估扁平字段）"],
            [
                "dispose",
                "处置执行",
                "预案自动触发或人工反制；未命中则解释原因",
                "done / current / skipped",
                "仅 done 且命中预案或人工处置时有；skipped 为空",
            ],
            [
                "result",
                "目标结果",
                "目标消失/迫降/打击/躁扰排除等闭环结论",
                "done（未闭环事件为 pending 且整节点隐藏）",
                "无（UI 不渲染阶段详情）",
            ],
        ],
        [900, 1100, 2400, 1500, 3460],
        font_size=8.8,
    )

    add_heading(document, "二、摘要文案与内置消息模板", 1)
    add_body(
        document,
        "各阶段摘要由 buildStageMessageInfo / buildAssessMessageInfo 生成，"
        "模板维护于 eventAttributeMessageAlign.ts（系统内置，无前台配置页）。"
        "占位符由事件字段填充。"
    )
    add_table(
        document,
        ["阶段", "事件名称", "消息模板（占位符）"],
        [
            ["目标发现", "目标发现", "{时间}{设备类型}{设备名称}{设备ID}发现目标"],
            ["威胁识别", "威胁识别", "{时间}根据多源数据融合成果，确认目标类型为{目标类型}"],
            ["威胁评估", "威胁评估", "{时间}{目标型号}{目标ID}触发威胁评估原则{规则名称}，判定为{等级}"],
            [
                "处置执行（自动）",
                "处置执行",
                "{时间}{目标型号}{目标ID}命中预案规则{规则名称}，执行{反制动作}，执行设备{设备名称}",
            ],
            [
                "处置执行（人工）",
                "处置执行",
                "{时间}{目标型号}{目标ID}经由人工处置，执行{反制动作}，执行设备{设备名称}",
            ],
            ["处置执行（未触发）", "处置执行", "{时间}未命中反制预案，未下发处置指令。"],
            [
                "目标结果（驱离/自离）",
                "目标结果",
                "{时间}，编码为{编码}的{无人机/非无人机}目标消失,逗留时间为{逗留时间}",
            ],
            ["目标结果（迫降）", "目标结果", "{时间}{目标型号}{目标ID}于{经纬度}迫降，请及时处理"],
            ["目标结果（打击）", "目标结果", "{时间}{目标型号}{目标ID}经反制打击，当前已无飞行特征信息"],
            ["目标结果（躁扰）", "目标结果", "{时间}事件已闭环，目标判定为躁扰信号排除。"],
        ],
        [1200, 1100, 7060],
        font_size=8.5,
    )

    add_heading(document, "三、威胁评估递进与 skipReason", 1)
    add_table(
        document,
        ["规则", "说明"],
        [
            ["递进路径", "最终高危 → [低危, 中危, 高危]；最终中危 → [低危, 中危]；最终低危 → [低危]；躁扰 → [无危]"],
            ["多条摘要", "递进时 threat 阶段使用 summaries 数组逐条展示；单步时使用 summary"],
            ["阶段详情", "每次评估扁平追加：评估次序（如有）、威胁等级、命中规则、评估依据、飞行状态"],
            ["白名单未反制", "skipReason：白名单目标未触发异常行为条件"],
            ["躁扰未反制", "skipReason：目标识别为飞鸟或躁扰信号，未进入反制处置流程"],
            ["低危未反制", "skipReason：威胁等级较低，未命中反制预案"],
            ["自然脱离", "skipReason：目标已自然离开防护区域（handlingResult 含自离/自然）"],
            ["默认未命中", "skipReason：未命中反制预案"],
        ],
        [2200, 7160],
        font_size=9.0,
    )

    add_heading(document, "四、场景识别条件对照", 1)
    add_table(
        document,
        ["场景", "关键识别条件", "目标类型 / 名单", "处置执行"],
        [
            ["黑飞无人机入侵", "uavSn=未解析 或 historyTargetType=黑飞", "黑飞无人机；名单多为未知", "视预案，通常可联动"],
            ["未知无人机入侵", "uavSn 可解析；名单=未知", "合作式无人机；未备案", "视威胁等级与预案"],
            ["黑名单无人机入侵", "名单=黑名单", "黑名单目标", "通常联动反制"],
            ["白名单无人机飞行", "名单=白名单；无异常行为", "白名单目标", "未触发（仅监测）"],
            ["躁扰告警", "historyTargetType=躁扰信号-飞鸟", "非无人机躁扰", "未触发"],
        ],
        [1600, 2800, 2400, 2560],
        font_size=8.8,
    )

    add_heading(document, "五、分场景触发流程与展示内容", 1)
    add_body(
        document,
        "以下各节描述该场景下会被触发并展示的内容。"
        "T1～T5 为阶段时间占位，实际时间由 discoveredAt 及演示偏移量推算；"
        "文案变量来自 HistoryEventItem 与预案/威胁规则命中结果。"
    )
    for scenario in SCENARIOS:
        add_scenario_section(
            document,
            scenario["title"],
            scenario["trigger"],
            scenario["flow"],
            scenario["stages"],
        )

    add_heading(document, "六、处置执行阶段详情字段（命中预案时）", 1)
    add_body(
        document,
        "当 disposalExecutionSource=plan 且 resolvePlanTriggerForEvent 命中时，"
        "buildDisposalExecutionStageDetails 输出以下字段；"
        "人工处置（manual）输出执行来源/操作人/选中设备/执行动作/预案策略=未自动触发；"
        "skipped 时 details 为空数组。"
    )
    add_table(
        document,
        ["字段", "来源", "说明"],
        [
            ["策略名称", "plan.planName", "命中预案名称"],
            ["触发规则", "triggerRule.ruleName", "预案内触发策略规则名"],
            ["天气要素", "triggerRule 天气条件", "温度/湿度/风力/降雨等，展示实际值与规则阈值"],
            ["所处区域", "triggerRule.areaLevel", "映射为区域名称列表"],
            ["反制动作", "deviceFunction / deviceAction", "如定向驱离、无线电压制、迫降处置"],
            ["联动设备", "countermeasureDevice", "去掉 (自动)/(人工) 等后缀后展示"],
        ],
        [1400, 2200, 5760],
        font_size=9.0,
    )

    add_heading(document, "七、跨场景差异摘要", 1)
    add_table(
        document,
        ["对比项", "黑飞", "未知", "黑名单", "白名单", "躁扰"],
        [
            ["识别码", "未解析", "可解析", "可解析", "可解析", "通常未解析/无关"],
            ["名单类型", "未知", "未知", "黑名单", "白名单", "未知"],
            ["威胁递进", "低→中→高（视最终等级）", "同左", "常到中/高危", "低危或无危", "仅无危"],
            ["处置执行", "可联动", "可联动", "联动", "未触发", "未触发"],
            ["目标结果模板", "驱离/自离/迫降/打击", "驱离/自离", "驱离/压制/迫降", "驱离/自离", "躁扰信号排除"],
            ["人工核查", "可选", "可选", "可选", "少见", "常见（确认躁扰）"],
        ],
        [1400, 1300, 1300, 1300, 1300, 1300],
        font_size=8.5,
    )

    add_heading(document, "八、与代码实现的对应关系", 1)
    add_bullet(document, "时间链构建：preview-app/src/api/lad/incident/disposalTimeline.ts → buildDisposalTimeline()")
    add_bullet(document, "摘要文案：preview-app/src/api/lad/incident/disposalTimelineMessage.ts")
    add_bullet(document, "预案命中详情：preview-app/src/api/lad/incident/planTriggerHitDetails.ts")
    add_bullet(document, "目标类型推断：preview-app/src/api/lad/incident/historyTargetType.ts")
    add_bullet(document, "UI 组件：preview-app/src/views/Lad/Incident/components/DisposalTimelinePanel.vue")
    add_bullet(document, "验证方式（自动识别/人工核查）体现在威胁识别阶段「识别结论」字段，不单独占阶段")
    add_bullet(document, "演示样本：躁扰场景可参考 mock 事件 he-10008（BIRD_NUISANCE_DEMO_EVENT_ID）")

    add_heading(document, "九、编写与验收要点", 1)
    add_bullet(document, "每个场景至少 1 条 mock 历史事件，覆盖五阶段或明确未触发原因。")
    add_bullet(document, "未触发处置执行时，摘要须可读，阶段详情为空，避免误以为系统故障。")
    add_bullet(document, "进行中事件（待处置/处置中）处置执行 status=current；目标结果 pending 时整节点不展示。")
    add_bullet(document, "躁扰类事件不得出现反制设备联动详情。")
    add_bullet(document, "目标结果阶段在 UI 中不可展开「阶段详情」。")
    add_bullet(document, "反制设备名、列表反制设备列均去掉 (自动)/(人工) 后缀。")

    return document


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    WORKSPACE_DOC.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)

    destinations = [WORKSPACE_DOC, PRODUCT_DOC, DESKTOP_DOC]
    saved = []
    for dest in destinations:
        dest.parent.mkdir(parents=True, exist_ok=True)
        try:
            shutil.copy2(OUTPUT_PATH, dest)
            saved.append(str(dest))
        except PermissionError:
            alt = dest.with_name(dest.stem + "-V2.docx")
            shutil.copy2(OUTPUT_PATH, alt)
            saved.append(str(alt))

    print(f"Generated: {OUTPUT_PATH}")
    for path in saved:
        print(f"Copied: {path}")


if __name__ == "__main__":
    main()
