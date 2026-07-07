from __future__ import annotations

import json
import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "系统三维地图功能设计与实现说明-工作交接版-V1.0.md"
DOCX = ROOT / "系统三维地图功能设计与实现说明-工作交接版-V1.0-正式版.docx"
CONTENT_WIDTH_DXA = 9360


def normalize(text: str) -> str:
    return re.sub(r"[\s`*_#|>\-—·]+", "", text)


with ZipFile(DOCX) as archive:
    assert archive.testzip() is None
    names = set(archive.namelist())
    assert "word/document.xml" in names
    assert "word/styles.xml" in names
    assert "word/numbering.xml" in names

doc = Document(DOCX)
source_text = SOURCE.read_text(encoding="utf-8")
doc_text = "\n".join(p.text for p in doc.paragraphs)
for table in doc.tables:
    doc_text += "\n" + "\n".join(cell.text for row in table.rows for cell in row.cells)

source_headings = [
    re.sub(r"^#{2,4}\s+", "", line).strip()
    for line in source_text.splitlines()
    if re.match(r"^#{2,4}\s+", line)
]
doc_headings = [
    p.text.strip()
    for p in doc.paragraphs
    if p.style.name in {"Heading 1", "Heading 2", "Heading 3"}
]
assert source_headings == doc_headings, (len(source_headings), len(doc_headings))

required_phrases = [
    "项目本地服务器",
    "国产化",
    "相对地面高度 AGL",
    "目标融合结果",
    "白名单仅作为目标属性",
    "躁扰目标",
    "距离目标最近",
    "不少于六个月",
    "高优先级区域默认覆盖低优先级区域",
    "业务数据输入契约",
    "场景交互输出契约",
    "系统状态反馈与提示",
    "工作交接基线",
]
for phrase in required_phrases:
    assert normalize(phrase) in normalize(doc_text), phrase

assert "```mermaid" not in doc_text
assert len(doc.tables) == 10

for table in doc.tables:
    tblw = table._tbl.tblPr.find(qn("w:tblW"))
    tblind = table._tbl.tblPr.find(qn("w:tblInd"))
    assert tblw is not None and int(tblw.get(qn("w:w"))) == CONTENT_WIDTH_DXA
    assert tblind is not None and int(tblind.get(qn("w:w"))) == 120
    grid = [int(x.get(qn("w:w"))) for x in table._tbl.tblGrid.findall(qn("w:gridCol"))]
    assert sum(grid) == CONTENT_WIDTH_DXA
    for row in table.rows:
        assert len(row.cells) == len(grid)
        for index, cell in enumerate(row.cells):
            tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            assert tcw is not None and int(tcw.get(qn("w:w"))) == grid[index]

section = doc.sections[0]
assert round(section.page_width.inches, 3) == 8.5
assert round(section.page_height.inches, 3) == 11.0
assert round(section.left_margin.inches, 3) == 1.0
assert round(section.header_distance.inches, 3) == 0.492
assert doc.styles["Normal"].font.size.pt == 11
assert doc.styles["Heading 1"].font.size.pt == 16
assert doc.styles["Heading 2"].font.size.pt == 13
assert doc.styles["Heading 3"].font.size.pt == 12

with ZipFile(DOCX) as archive:
    document_xml = archive.read("word/document.xml").decode("utf-8")
    numbering_xml = archive.read("word/numbering.xml").decode("utf-8")
    assert document_xml.count("<pic:pic") == 2
    assert document_xml.count("w:numPr") >= 20
    assert "三维地图模块关系与业务闭环" in document_xml
    assert "实时监测与快速处置流程" in document_xml
    assert "w:abstractNum" in numbering_xml and "w:numFmt" in numbering_xml

report = {
    "docx": str(DOCX),
    "headings": len(doc_headings),
    "paragraphs": len(doc.paragraphs),
    "tables": len(doc.tables),
    "figures": 2,
    "page_geometry": "Letter portrait, 1-inch margins",
    "table_geometry_dxa": CONTENT_WIDTH_DXA,
    "content_checks": len(required_phrases),
    "result": "PASS",
}
print(json.dumps(report, ensure_ascii=False, indent=2))
